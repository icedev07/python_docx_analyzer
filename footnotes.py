from docx import Document
from docx.oxml.ns import qn
from xml.etree.ElementTree import fromstring

def get_footnotes_part(doc):
    for rel in doc.part.rels.values():
        if "footnotes" in rel.target_ref:
            return rel.target_part
    return None

def extract_footnotes(doc):
    footnotes_part = get_footnotes_part(doc)
    footnotes = {}
    if footnotes_part:
        footnotes_xml = fromstring(footnotes_part.blob)
        for footnote in footnotes_xml.findall(qn('w:footnote')):
            footnote_id = int(footnote.get(qn('w:id')))
            footnote_text = ''
            for paragraph in footnote.findall(qn('w:p')):
                for run in paragraph.findall(qn('w:r')):
                    text = ''.join([node.text for node in run.findall(qn('w:t')) if node.text])
                    footnote_text += text
            footnotes[footnote_id] = footnote_text
    return footnotes

def copy_footnotes(src_doc, dest_doc, footnotes):
    for paragraph in src_doc.paragraphs:
        for run in paragraph.runs:
            for element in run._element:
                if element.tag == qn('w:footnoteReference'):
                    footnote_id = int(element.get(qn('w:id')))
                    # Add the text of the footnote to the destination document
                    footnote_text = footnotes.get(footnote_id, "")
                    dest_doc.add_paragraph(f"Footnote {footnote_id}: {footnote_text}")

# Load the source document
src_doc = Document('input.docx')

# Create a new destination document
dest_doc = Document()

# Extract footnotes from the source document
footnotes = extract_footnotes(src_doc)

# Copy footnotes to the destination document
copy_footnotes(src_doc, dest_doc, footnotes)

# Save the destination document
dest_doc.save('destination.docx')

print("Footnotes have been successfully copied to destination.docx.")
