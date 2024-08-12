from docx import Document
from docx.oxml import OxmlElement, ns
import os

# Load the original document
doc = Document('input.docx')

# Create a new document
new_doc = Document()

# Directory to save images (if there are any)
image_dir = "extracted_images"
os.makedirs(image_dir, exist_ok=True)

# Footnote map to keep track of footnotes
footnote_map = {}

# Function to add a footnote reference to a paragraph
def add_footnote_reference(paragraph, footnote_id):
    run = paragraph.add_run()
    footnote_ref = OxmlElement('w:footnoteReference')
    footnote_ref.set(ns.qn('w:id'), str(footnote_id))
    run._r.append(footnote_ref)

# Function to copy the runs and text, including handling footnotes
def copy_runs(source_paragraph, target_paragraph):
    for run in source_paragraph.runs:
        new_run = target_paragraph.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.name = run.font.name
        new_run.font.size = run.font.size
        if run.font.color:
            new_run.font.color.rgb = run.font.color.rgb

        # Handle footnotes
        for element in run._element:
            if element.tag.endswith('footnoteReference'):
                footnote_id = element.get(ns.qn('w:id'))
                if footnote_id not in footnote_map:
                    footnote_map[footnote_id] = f"[{footnote_id}] {run.text}"
                add_footnote_reference(target_paragraph, footnote_id)

# Function to copy paragraphs, handling runs and footnotes
def copy_paragraph(source_paragraph, target_paragraph):
    copy_runs(source_paragraph, target_paragraph)
    target_paragraph.alignment = source_paragraph.alignment
    target_paragraph.paragraph_format.left_indent = source_paragraph.paragraph_format.left_indent
    target_paragraph.paragraph_format.right_indent = source_paragraph.paragraph_format.right_indent
    target_paragraph.paragraph_format.space_before = source_paragraph.paragraph_format.space_before
    target_paragraph.paragraph_format.space_after = source_paragraph.paragraph_format.space_after
    target_paragraph.paragraph_format.line_spacing = source_paragraph.paragraph_format.line_spacing
    target_paragraph.paragraph_format.keep_together = source_paragraph.paragraph_format.keep_together
    target_paragraph.paragraph_format.keep_with_next = source_paragraph.paragraph_format.keep_with_next
    target_paragraph.paragraph_format.page_break_before = source_paragraph.paragraph_format.page_break_before
    target_paragraph.paragraph_format.widow_control = source_paragraph.paragraph_format.widow_control

# Copy all the paragraphs and handle footnotes
for paragraph in doc.paragraphs:
    new_paragraph = new_doc.add_paragraph()
    copy_paragraph(paragraph, new_paragraph)

# Copy headers and footers
for section in doc.sections:
    new_section = new_doc.sections[0]

    # Copy header
    for paragraph in section.header.paragraphs:
        new_paragraph = new_section.header.add_paragraph()
        copy_paragraph(paragraph, new_paragraph)

    # Copy footer
    for paragraph in section.footer.paragraphs:
        new_paragraph = new_section.footer.add_paragraph()
        copy_paragraph(paragraph, new_paragraph)

# Add footnotes to the end of the document
if footnote_map:
    new_doc.add_paragraph("\nFootnotes")
    for footnote_id, footnote_text in footnote_map.items():
        footnote_paragraph = new_doc.add_paragraph()
        footnote_paragraph.add_run(footnote_text)

# Save the new document
new_doc.save('result.docx')
