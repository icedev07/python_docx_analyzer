import os
import zipfile
import xml.etree.ElementTree as ET
import shutil
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Function to extract a .docx file
def extract_docx(docx_path, extract_dir):
    with zipfile.ZipFile(docx_path, 'r') as zip_ref:
        zip_ref.extractall(extract_dir)

# Function to recreate a .docx file from the extracted content
def recreate_docx(original_extract_dir, new_docx_path):
    with zipfile.ZipFile(new_docx_path, 'w') as docx:
        for foldername, subfolders, filenames in os.walk(original_extract_dir):
            for filename in filenames:
                file_path = os.path.join(foldername, filename)
                arcname = os.path.relpath(file_path, original_extract_dir)
                docx.write(file_path, arcname)

# Translation function (replace this with a real translation implementation)
def translate_text(text, target_language='en'):
    # Dummy translation function, replace with an actual translation service like Google Translate API
    return text + '++'

# Function to translate text while preserving all styles and formatting
def translate_paragraph_with_formatting(paragraph, target_language='en'):

    for run in paragraph.runs:
        if run.text.strip():
            original_text = run.text
            # translated_text = translator.translate(original_text, dest=target_language).text
            translated_text = translate_text(original_text)
            run.text = translated_text  # Replace text while keeping the style intact

# Function to translate headers, footers, and footnotes if they exist
def translate_headers_footers_footnotes(doc, extract_dir, target_language='en'):
    # Translate headers and footers
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                translate_paragraph_with_formatting(para, target_language)

        if section.footer:
            for para in section.footer.paragraphs:
                translate_paragraph_with_formatting(para, target_language)

    # Translate footnotes (as an example, handling XML directly here)
    footnotes_file = os.path.join(extract_dir, 'word/footnotes.xml')
    if os.path.exists(footnotes_file):
        with open(footnotes_file, 'r', encoding='utf-8') as file:
            footnotes_content = file.read()

        # You can apply similar logic to translate the content while preserving footnote formatting

# Main function to handle translation of a .docx file
def analyze_and_translate_docx(input_docx, output_docx, target_language='en'):
    # Extract the .docx file content
    extract_dir = 'extracted_docx'
    if os.path.exists(extract_dir):
        shutil.rmtree(extract_dir)
    os.makedirs(extract_dir)
    
    # Extract content
    extract_docx(input_docx, extract_dir)

    # Load the original document
    doc = Document(input_docx)

    # Translate document body while preserving formatting
    for para in doc.paragraphs:
        translate_paragraph_with_formatting(para, target_language)

    # Translate headers, footers, and footnotes
    translate_headers_footers_footnotes(doc, extract_dir, target_language)

    # Save the modified document
    doc.save(output_docx)

    # Clean up extracted files
    shutil.rmtree(extract_dir)

# Example usage:
input_docx = 'input.docx'  # Path to your input .docx file
output_docx = 'output.docx'  # Path to the output .docx file to be generated
target_language = 'en'  # Target language code (e.g., 'en' for English, 'fr' for French)
analyze_and_translate_docx(input_docx, output_docx, target_language)
